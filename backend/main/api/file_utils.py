# api/file_utils.py
"""
Funktionen zur Verarbeitung und Extraktion von Dateien verschiedener Formate.
"""

import io
import logging
import traceback

import fitz  # PyMuPDF

from .error_handler import log_error

logger = logging.getLogger(__name__)


def allowed_file(filename):
    """
    Überprüft, ob der Dateityp zulässig ist.

    Args:
        filename (str): Name der Datei

    Returns:
        bool: True, wenn der Dateityp zulässig ist, sonst False
    """
    allowed_extensions = {'pdf', 'txt', 'docx', 'doc'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions


def extract_text_from_file(file):
    """
    Extrahiert Text aus einer Datei basierend auf ihrem Typ.

    Args:
        file: Ein Datei-Objekt (z.B. aus request.files)

    Returns:
        tuple: (extrahierter Text, Dateityp, Dateigröße, Seitenanzahl)
    """
    # Versuche, den Dateiinhalt zu lesen
    try:
        filename = file.filename
        file_content = file.read()

        # Überprüfe Dateigröße
        file_size = len(file_content)
        max_size = 20 * 1024 * 1024  # 20 MB

        if file_size > max_size:
            logger.warning("Datei zu groß: %s (%s MB)", filename, f"{file_size/(1024*1024):.2f}")
            return None, None, file_size, 0

        # Setze Dateiposition zurück für eventuell weitere Operationen
        file.seek(0)

        # Extrahiere Text basierend auf Dateityp
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        text = ""
        page_count = 0

        if file_ext == 'pdf':
            text, page_count = extract_text_from_pdf(file_content)

        elif file_ext == 'txt':
            try:
                text = file_content.decode('utf-8', errors='replace')
                page_count = 1  # Txt hat immer nur eine "Seite"
            except Exception as e:
                logger.error("Fehler beim Dekodieren der TXT-Datei: %s", str(e))
                return None, file_ext, file_size, 0

        elif file_ext in ['docx', 'doc']:
            from io import BytesIO
            try:
                logger.info("Verarbeite DOCX/DOC: %s", filename)
                import docx
                doc = docx.Document(BytesIO(file_content))
                text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                page_count = len(doc.sections)  # Approximation
            except Exception as docx_err:
                # Fallback für alte .doc-Dateien
                try:
                    logger.info("DOCX-Verarbeitung fehlgeschlagen, versuche DOC-Fallback für %s", filename)
                    try:
                        from antiword import Document
                        doc = Document(file_content)
                        text = doc.getText()
                        page_count = 1  # Standardannahme
                    except ImportError:
                        logger.warning("antiword-Bibliothek ist nicht installiert, überspringe DOC-Verarbeitung")
                        raise  # Weiterleiten an äußeren Exception-Handler
                except Exception as doc_err:
                    logger.error("Fehler bei DOC/DOCX-Verarbeitung: %s / %s", str(docx_err), str(doc_err))
                    return None, file_ext, file_size, 0
        else:
            logger.warning("Nicht unterstützter Dateityp: %s", file_ext)
            return None, file_ext, file_size, 0

        # Prüfe, ob Text extrahiert wurde
        if not text or not text.strip():
            logger.warning("Kein Text aus %s extrahiert", filename)
            return None, file_ext, file_size, page_count

        # Gib extrahierten Text, Dateityp, Größe und Seitenanzahl zurück
        return text, file_ext, file_size, page_count

    except Exception as e:
        # Allgemeine Fehlerbehandlung mit Stacktrace
        stack_trace = traceback.format_exc()
        logger.error("Fehler beim Extrahieren von Text: %s\n%s", str(e), stack_trace)
        return None, None, 0, 0


def extract_text_from_pdf(file_data):
    """
    Extrahiert Text aus einer PDF-Datei mit PyMuPDF.

    Args:
        file_data (bytes): Binärdaten der PDF-Datei

    Returns:
        tuple: (extrahierter Text, Seitenanzahl)
    """
    try:
        # Sammle extrahierten Text
        all_text = []
        page_count = 0

        # Öffne das PDF als Speicherobjekt
        with io.BytesIO(file_data) as data:
            try:
                # Öffne das PDF-Dokument
                pdf_document = fitz.open(stream=data, filetype="pdf")
                page_count = len(pdf_document)

                # Extrahiere Text von jeder Seite
                for page_num in range(page_count):
                    try:
                        page = pdf_document.load_page(page_num)
                        text = page.get_text("text")
                        if text:
                            all_text.append(text)
                    except Exception as page_error:
                        logger.warning("Fehler bei der Extraktion der Seite %s: %s", page_num+1, str(page_error))
                        continue

                # Schließe das Dokument
                pdf_document.close()
            except Exception as e:
                logger.warning("Fehler bei der PDF-Extraktion: %s", str(e))
                return None, page_count

        # Kombiniere den gesamten Text
        final_text = "\n\n".join([text for text in all_text if text.strip()])

        # Wenn kein Text extrahiert werden konnte
        if not final_text.strip():
            return (
                "Der Text konnte aus dieser PDF nicht extrahiert werden. "
                "Es könnte sich um eine gescannte PDF ohne OCR handeln."
            ), page_count

        # Minimale Bereinigung - nur NUL-Bytes entfernen
        final_text = final_text.replace('\x00', '')

        return final_text, page_count
    except Exception as e:
        logger.error("Kritischer Fehler bei der PDF-Extraktion: %s", str(e))
        return None, 0


def extract_text_from_pdf_safe(file_data):
    """
    Sichere Version der PDF-Textextraktion mit erweiterten Fehlerbehandlungen.

    Args:
        file_data (bytes): Binärdaten der PDF-Datei

    Returns:
        str: Extrahierter Text oder Fehlermeldung
    """
    try:
        # Sammle extrahierten Text
        all_text = []

        # Öffne das PDF als Speicherobjekt
        with io.BytesIO(file_data) as data:
            try:
                # Öffne das PDF-Dokument ohne Änderungen an den Binärdaten
                pdf_document = fitz.open(stream=data, filetype="pdf")

                # Extrahiere Text von jeder Seite mit sanften Einstellungen
                for page_num in range(len(pdf_document)):
                    try:
                        page = pdf_document.load_page(page_num)
                        # Verwende einfache Textextraktion ohne Formatierung oder Bereinigung
                        text = page.get_text("text")
                        if text:
                            all_text.append(text)
                    except Exception as page_error:
                        logger.warning(
                            f"Fehler bei der Extraktion der Seite {page_num+1} mit PyMuPDF: {str(page_error)}")
                        continue

                # Schließe das Dokument
                pdf_document.close()
            except Exception as e:
                logger.warning("Fehler bei der Extraktion mit PyMuPDF: %s", str(e))
                return f"CORRUPTED_PDF: {str(e)}"

        # Kombiniere den gesamten Text
        final_text = "\n\n".join([text for text in all_text if text.strip()])

        # Wenn noch immer kein Text, gib einen klaren Hinweis zurück
        if not final_text.strip():
            return (
                "Der Text konnte aus dieser PDF nicht extrahiert werden. "
                "Es könnte sich um eine gescannte PDF ohne OCR, eine beschädigte Datei "
                "oder eine stark gesicherte PDF handeln."
            )

        # Minimale Bereinigung - nur NUL-Bytes entfernen, da diese in Datenbanken Probleme verursachen können
        final_text = final_text.replace('\x00', '')

        return final_text
    except Exception as e:
        logger.error("Kritischer Fehler bei PDF-Extraktionsversuch mit PyMuPDF: %s", str(e))
        return f"CORRUPTED_PDF: {str(e)}"
