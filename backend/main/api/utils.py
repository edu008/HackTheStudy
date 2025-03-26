import json
import re
import docx
from langdetect import detect
import io  # Hinzugefügt für BytesIO
from flask import current_app, g, jsonify, request
import logging
import fitz  # PyMuPDF
import time  # Für sleep-Funktion
import os
import sys
from core.models import db, Upload, Flashcard, Question, Topic, Connection, UserActivity
from core.redis_client import redis_client
import tiktoken
import random
from functools import lru_cache
import hashlib
from datetime import datetime
from openaicache.openai_wrapper import CachedOpenAI
from openaicache.token_tracker import TokenTracker
from . import api_bp
from .error_handler import (
    log_error, create_error_response, ERROR_INVALID_INPUT, ERROR_FILE_PROCESSING,
    InvalidInputError, FileProcessingError
)

# Neue Importe für verbesserte Funktionalität
from api.log_utils import AppLogger
from api.openai_client import OptimizedOpenAIClient

# Re-Export wichtiger Funktionen aus token_tracking für Abwärtskompatibilität
from api.token_tracking import count_tokens, calculate_token_cost, check_credits_available, deduct_credits

logger = logging.getLogger(__name__)

def allowed_file(filename):
    allowed_extensions = {'pdf', 'txt', 'docx', 'doc'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions

def clean_text_for_database(text):
    """
    Bereinigt einen Text, um sicherzustellen, dass er in der Datenbank gespeichert werden kann.
    Entfernt NUL-Zeichen (0x00) und andere problematische Zeichen.
    
    Args:
        text (str): Der zu bereinigende Text
        
    Returns:
        str: Der bereinigte Text
    """
    if not text:
        return ""
    
    try:
        # Entferne Null-Bytes (0x00)
        cleaned_text = text.replace('\x00', '')
        
        # Aggressivere Bereinigung - alle Steuerzeichen ausser Zeilenumbrüche und Tabs entfernen
        # Dies verhindert viele Probleme mit exotischen PDF-Formaten
        allowed_control = ['\n', '\r', '\t']
        cleaned_text = ''.join(c for c in cleaned_text if c >= ' ' or c in allowed_control)
        
        # Bereinige Unicode-Escape-Sequenzen, die Probleme verursachen könnten
        cleaned_text = cleaned_text.encode('utf-8', errors='ignore').decode('utf-8', errors='ignore')
        
        # Entferne übermässige Leerzeichen und Zeilenumbrüche
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)  # Mehr als 2 aufeinanderfolgende Zeilenumbrüche reduzieren
        cleaned_text = re.sub(r' {3,}', '  ', cleaned_text)     # Mehr als 2 aufeinanderfolgende Leerzeichen reduzieren
        
        return cleaned_text
    except Exception as e:
        log_error(e, endpoint="clean_text_for_database")
        # Im Fehlerfall einen sicheren leeren String zurückgeben
        return ""

def extract_text_from_file(file_content, filename):
    """
    Extrahiert Text aus verschiedenen Dateiformaten.
    
    Args:
        file_content: Der Inhalt der Datei als Bytes oder Hex-String
        filename: Der Name der Datei zur Bestimmung des Typs
        
    Returns:
        str: Der extrahierte Text
    """
    import logging
    import traceback
    logger = logging.getLogger(__name__)
    
    try:
        # Konvertiere hex zu bytes falls nötig
        if isinstance(file_content, str):
            try:
                file_content = bytes.fromhex(file_content)
            except ValueError as e:
                logger.error(f"Fehler beim Konvertieren des Hex-String: {str(e)}")
                return f"ERROR: Konnte Datei nicht verarbeiten - {str(e)}"
        
        if not isinstance(file_content, bytes):
            return "ERROR: Dateityp nicht unterstützt oder kein Inhalt"
        
        # Prüfe auf leere Datei
        if len(file_content) == 0:
            logger.warning(f"Leere Datei: {filename}")
            return "ERROR: Die Datei ist leer"
            
        # Begrenze die Größe
        max_file_size = 20 * 1024 * 1024  # 20 MB
        if len(file_content) > max_file_size:
            logger.warning(f"Datei zu groß: {filename} ({len(file_content) // (1024*1024)} MB)")
            return f"ERROR: Die Datei ist zu groß ({len(file_content) // (1024*1024)} MB)"
        
        file_ext = filename.lower().split('.')[-1] if '.' in filename else ''
        
        # Ausführliche Protokollierung für alle Operationen
        logger.info(f"Verarbeite Datei: {filename} (Typ: {file_ext}, Größe: {len(file_content)} Bytes)")
        
        if file_ext == 'pdf':
            # Verwende nur PyMuPDF für PDF-Extraktion
            try:
                return extract_text_from_pdf_safe(file_content)
            except Exception as e:
                logger.error(f"Fehler bei der PDF-Extraktion mit PyMuPDF: {str(e)}")
                return f"ERROR: Fehler bei der PDF-Extraktion: {str(e)}"
                
        elif file_ext == 'txt':
            try:
                return file_content.decode('utf-8', errors='replace')
            except Exception as e:
                logger.error(f"Fehler beim Dekodieren der TXT-Datei: {str(e)}")
                return f"ERROR: Textdatei konnte nicht dekodiert werden: {str(e)}"
                
        elif file_ext in ['docx', 'doc']:
            from io import BytesIO
            try:
                logger.info(f"Verarbeite DOCX/DOC: {filename}")
                import docx
                doc = docx.Document(BytesIO(file_content))
                return "\n".join([paragraph.text for paragraph in doc.paragraphs])
            except Exception as docx_err:
                # Fallback für alte .doc-Dateien
                try:
                    logger.info(f"DOCX-Verarbeitung fehlgeschlagen, versuche DOC-Fallback für {filename}")
                    from antiword import Document
                    doc = Document(file_content)
                    return doc.getText()
                except Exception as doc_err:
                    logger.error(f"Fehler bei DOC/DOCX-Verarbeitung: {str(docx_err)} / {str(doc_err)}")
                    return f"ERROR: Dokument konnte nicht gelesen werden: {str(doc_err)}"
        else:
            logger.warning(f"Nicht unterstützter Dateityp: {file_ext}")
            return f"ERROR: Dateityp .{file_ext} wird nicht unterstützt."
    
    except RuntimeError as rt_err:
        # Detaillierte Protokollierung für RuntimeErrors
        stack_trace = traceback.format_exc()
        logger.critical(f"Kritischer RuntimeError bei der Textextraktion: {str(rt_err)}\n{stack_trace}")
        return f"CRITICAL_ERROR: RuntimeError bei der Textextraktion: {str(rt_err)}"
        
    except Exception as e:
        # Allgemeine Fehlerbehandlung mit Stacktrace
        stack_trace = traceback.format_exc()
        logger.error(f"Fehler beim Extrahieren von Text aus {filename}: {str(e)}\n{stack_trace}")
        return f"ERROR: Fehler beim Lesen der Datei: {str(e)}"

def extract_text_from_pdf_safe(file_data):
    """
    Extrahiert Text aus einer PDF-Datei mit PyMuPDF (fitz).
    PyMuPDF kann robuster mit verschiedenen PDF-Formaten umgehen.
    """
    try:
        # Sammle extrahierten Text
        all_text = []
        
        # Öffne das PDF als Speicherobjekt
        with io.BytesIO(file_data) as data:
            try:
                # Öffne das PDF-Dokument ohne Änderungen an den Binärdaten
                pdf_document = fitz.open(stream=data, filetype="pdf")
                
                # Extrahiere Text von jeder Seite mit sanften Einstellungen
                for page_num in range(len(pdf_document)):
                    try:
                        page = pdf_document.load_page(page_num)
                        # Verwende einfache Textextraktion ohne Formatierung oder Bereinigung
                        text = page.get_text("text")
                        if text:
                            all_text.append(text)
                    except Exception as page_error:
                        logger.warning(f"Fehler bei der Extraktion der Seite {page_num+1} mit PyMuPDF: {str(page_error)}")
                        continue
                
                # Schliesse das Dokument
                pdf_document.close()
            except Exception as e:
                logger.warning(f"Fehler bei der Extraktion mit PyMuPDF: {str(e)}")
                return f"CORRUPTED_PDF: {str(e)}"
        
        # Kombiniere den gesamten Text
        final_text = "\n\n".join([text for text in all_text if text.strip()])
        
        # Wenn noch immer kein Text, gib einen klaren Hinweis zurück
        if not final_text.strip():
            return "Der Text konnte aus dieser PDF nicht extrahiert werden. Es könnte sich um eine gescannte PDF ohne OCR, eine beschädigte Datei oder eine stark gesicherte PDF handeln."
        
        # Minimale Bereinigung - nur NUL-Bytes entfernen, da diese in Datenbanken Probleme verursachen können
        final_text = final_text.replace('\x00', '')
        
        return final_text
    except Exception as e:
        logger.error(f"Kritischer Fehler bei PDF-Extraktionsversuch mit PyMuPDF: {str(e)}")
        return f"CORRUPTED_PDF: {str(e)}"

def extract_text_from_pdf(file_data):
    """Extrahiert Text aus einer PDF-Datei mit PyMuPDF."""
    try:
        # Verwende PyMuPDF (fitz) ohne Bereinigung der Binärdaten
        text = extract_text_from_pdf_safe(file_data)
        
        # Prüfe, ob die Methode einen CORRUPTED_PDF Fehlercode zurückgegeben hat
        if isinstance(text, str) and text.startswith('CORRUPTED_PDF:'):
            return text
            
        if text and text.strip():
            return text
        else:
            return "Keine Textdaten konnten aus dieser PDF-Datei extrahiert werden. Es könnte sich um ein Scan-Dokument ohne OCR handeln."
    except Exception as e:
        logger.error(f"Kritischer Fehler bei der PDF-Extraktion: {str(e)}")
        return f"CORRUPTED_PDF: {str(e)}"

def detect_language(text):
    try:
        lang = detect(text[:500])  # Begrenze auf 500 Zeichen für Effizienz
        return 'de' if lang == 'de' else 'en'
    except Exception:
        return 'en'

# Cache für OpenAI-Antworten (Kann später entfernt werden)
_response_cache = {}

def query_chatgpt(prompt, client, system_content=None, temperature=0.7, max_retries=5, use_cache=True, session_id=None, function_name="query_chatgpt"):
    """
    Sendet eine Anfrage an die OpenAI API mit Caching und Token-Tracking.
    
    Args:
        prompt: Der Prompt-Text für die Anfrage
        client: Der OpenAI-Client (wird nur für Legacy-Kompatibilität verwendet)
        system_content: Optionaler System-Prompt
        temperature: Temperatur für die Antwortgenerierung (0.0-1.0)
        max_retries: Maximale Anzahl von Wiederholungsversuchen
        use_cache: Ob der Cache verwendet werden soll (default: True)
        session_id: ID der aktuellen Session (für Token-Tracking)
        function_name: Name der aufrufenden Funktion (für Token-Tracking)
        
    Returns:
        Die Antwort der API oder eine Fehlermeldung
    """
    openai_logger = logging.getLogger('openai_api')
    openai_logger.setLevel(logging.INFO)
    
    # Logge den Prompt
    prompt_preview = prompt[:500] + "..." if len(prompt) > 500 else prompt
    openai_logger.info(f"🔵 OPENAI ANFRAGE [{function_name}] - Session: {session_id or 'unbekannt'}")
    openai_logger.info(f"📝 PROMPT: {prompt_preview}")
    if system_content:
        system_preview = system_content[:300] + "..." if len(system_content) > 300 else system_content
        openai_logger.info(f"⚙️ SYSTEM: {system_preview}")
    openai_logger.info(f"🧮 PARAMETER: model=?, temperature={temperature}, use_cache={use_cache}")
    
    request_start_time = time.time()
    
    try:
        # Import hier, um Zirkelbezüge zu vermeiden
        from openaicache import get_openai_client
        
        # Verwende unser neues Caching-System
        cached_client = get_openai_client(use_cache=use_cache)
        
        # Modell aus der Konfiguration abrufen
        try:
            model = current_app.config.get('OPENAI_MODEL', 'gpt-4o').strip()
        except RuntimeError:
            model = os.getenv('OPENAI_MODEL', 'gpt-4o').strip()
        
        openai_logger.info(f"🤖 MODELL: {model}")
        
        # Bestimme die aktuelle Benutzer-ID
        user_id = None
        if hasattr(g, 'user') and g.user:
            user_id = g.user.id
        
        # Sende Anfrage mit dem Cached-Client
        response = cached_client.chat_completion(
            prompt=prompt,
            system_content=system_content,
            model=model,
            temperature=temperature,
            max_tokens=4000,
            use_cache=use_cache,
            user_id=user_id,
            session_id=session_id,
            function_name=function_name,
            endpoint=function_name,
            max_retries=max_retries
        )
        
        request_time = time.time() - request_start_time
        
        # Fehlerbehandlung
        if 'error' in response:
            # Spezialfall "insufficient_credits" - direkt zurückgeben
            if response.get('error_type') == 'insufficient_credits':
                error_msg = response['error']
                openai_logger.error(f"❌ OPENAI FEHLER [{function_name}]: {error_msg} - Zeit: {request_time:.2f}s")
                raise ValueError(error_msg)
            
            # Andere Fehler weiterleiten
            error_msg = f"OpenAI API Error: {response.get('error')}"
            openai_logger.error(f"❌ OPENAI FEHLER [{function_name}]: {error_msg} - Zeit: {request_time:.2f}s")
            raise ValueError(error_msg)
        
        # Logge die Antwort
        answer_text = response.get('text', 'Keine Antwort')
        answer_preview = answer_text[:500] + "..." if len(answer_text) > 500 else answer_text
        
        # Token-Informationen
        token_info = response.get('token_info', {})
        input_tokens = token_info.get('input_tokens', 0)
        output_tokens = token_info.get('output_tokens', 0)
        total_tokens = input_tokens + output_tokens
        cost = token_info.get('cost', 0)
        
        openai_logger.info(f"🟢 OPENAI ANTWORT [{function_name}] - Zeit: {request_time:.2f}s - Tokens: {input_tokens}/{output_tokens} (Kosten: {cost})")
        openai_logger.info(f"📄 ANTWORT: {answer_preview}")
        
        # Gib den Antworttext zurück
        return response['text']
    
    except ImportError as ie:
        # Fallback auf den ursprünglichen Weg, wenn das Cache-System nicht verfügbar ist
        logger.warning(f"Cache-System nicht verfügbar: {str(ie)}. Verwende Fallback-Implementierung.")
        openai_logger.warning(f"⚠️ FALLBACK AUF LEGACY-IMPLEMENTIERUNG: {str(ie)}")
        return _query_chatgpt_legacy(prompt, client, system_content, temperature, max_retries, use_cache, session_id, function_name)
    
    except Exception as e:
        request_time = time.time() - request_start_time
        error_msg = f"OpenAI API Error: {str(e)}"
        logger.error(f"Fehler bei der OpenAI-Anfrage: {str(e)}")
        openai_logger.error(f"❌ OPENAI FEHLER [{function_name}]: {str(e)} - Zeit: {request_time:.2f}s")
        raise ValueError(error_msg)

def _query_chatgpt_legacy(prompt, client, system_content=None, temperature=0.7, max_retries=5, use_cache=True, session_id=None, function_name="query_chatgpt"):
    """
    Ursprüngliche Implementierung der OpenAI-Anfrage für Fallback-Zwecke.
    """
    openai_logger = logging.getLogger('openai_api')
    openai_logger.setLevel(logging.INFO)
    
    # Logge den Prompt (Legacy)
    prompt_preview = prompt[:500] + "..." if len(prompt) > 500 else prompt
    openai_logger.info(f"🔵 OPENAI LEGACY ANFRAGE [{function_name}] - Session: {session_id or 'unbekannt'}")
    openai_logger.info(f"📝 LEGACY PROMPT: {prompt_preview}")
    if system_content:
        system_preview = system_content[:300] + "..." if len(system_content) > 300 else system_content
        openai_logger.info(f"⚙️ LEGACY SYSTEM: {system_preview}")
    
    request_start_time = time.time()
    
    try:
        # Wähle das Modell basierend auf Konfiguration oder Umgebungsvariablen
        try:
            model = current_app.config.get('OPENAI_MODEL', 'gpt-4o')
        except RuntimeError:
            model = os.getenv('OPENAI_MODEL', 'gpt-4o')
        
        openai_logger.info(f"🤖 LEGACY MODELL: {model}")
        
        # Cache-Schlüssel generieren, falls Caching aktiviert ist
        cache_key = None
        if use_cache:
            key_parts = [prompt, model, str(temperature)]
            if system_content:
                key_parts.append(system_content)
                
            cache_key = hashlib.md5('_'.join(key_parts).encode()).hexdigest()
            
            # Prüfe, ob das Ergebnis im Cache ist
            if cache_key in _response_cache:
                cached_result = _response_cache[cache_key]
                openai_logger.info(f"🔄 CACHE-TREFFER für '{function_name}' - Verwende gecachte Antwort")
                return cached_result
        
        # Bereite Nachrichten vor
        messages = []
        
        # Füge System-Message hinzu, falls vorhanden
        if system_content:
            messages.append({
                "role": "system",
                "content": system_content
            })
        
        # Füge User-Message hinzu
        messages.append({
            "role": "user",
            "content": prompt
        })
        
        # Führe die API-Anfrage durch
        try:
            # Import für Token-Tracking
            from api.token_tracking import track_token_usage, check_credits_available
            
            # Prüfe, ob genügend Credits vorhanden sind
            user_id = None
            if hasattr(g, 'user') and g.user:
                user_id = g.user.id
                
            if user_id and session_id:
                # Prüfe Credits für den Benutzer
                credits_check = check_credits_available(user_id, prompt, model)
                if not credits_check["success"]:
                    openai_logger.error(f"❌ LEGACY CREDITS FEHLER: {credits_check['message']}")
                    raise ValueError(credits_check["message"])
            
            # Sende Anfrage an OpenAI
            response = client.chat.completions.create(
                model=model,
                messages=messages,
                temperature=temperature,
                max_tokens=4000
            )
            
            # Extrahiere Antworttext
            if response and response.choices and len(response.choices) > 0:
                result = response.choices[0].message.content
                
                # Speichere im Cache, falls aktiviert
                if use_cache and cache_key:
                    _response_cache[cache_key] = result
                
                # Token-Tracking
                if user_id and session_id:
                    track_token_usage(
                        user_id=user_id,
                        input_text=prompt,
                        output_text=result,
                        session_id=session_id,
                        function_name=function_name,
                        input_tokens=response.usage.prompt_tokens,
                        output_tokens=response.usage.completion_tokens
                    )
                
                # Logge die Antwort
                request_time = time.time() - request_start_time
                result_preview = result[:500] + "..." if len(result) > 500 else result
                
                # Log Token-Informationen
                input_tokens = response.usage.prompt_tokens
                output_tokens = response.usage.completion_tokens
                total_tokens = input_tokens + output_tokens
                
                openai_logger.info(f"🟢 LEGACY ANTWORT [{function_name}] - Zeit: {request_time:.2f}s - Tokens: {input_tokens}/{output_tokens}")
                openai_logger.info(f"📄 LEGACY ANTWORT: {result_preview}")
                
                return result
            else:
                error_msg = "OpenAI API gab keine gültige Antwort zurück"
                openai_logger.error(f"❌ LEGACY FEHLER: {error_msg}")
                raise ValueError(error_msg)
                
        except Exception as api_error:
            request_time = time.time() - request_start_time
            error_msg = f"OpenAI API Error: {str(api_error)}"
            openai_logger.error(f"❌ LEGACY API FEHLER [{function_name}]: {str(api_error)} - Zeit: {request_time:.2f}s")
            raise ValueError(error_msg)
    
    except Exception as outer_error:
        request_time = time.time() - request_start_time
        error_msg = f"OpenAI Legacy API Error: {str(outer_error)}"
        openai_logger.error(f"❌ LEGACY GESAMTFEHLER [{function_name}]: {str(outer_error)} - Zeit: {request_time:.2f}s")
        raise ValueError(error_msg)

def analyze_content(text, client, language='en', session_id=None, function_name="analyze_content"):
    system_content = (
        "You are an expert in educational content analysis with a specialization in creating hierarchical knowledge structures."
        if language != 'de' else
        "Sie sind ein Experte für die Analyse von Bildungsinhalten mit einer Spezialisierung auf die Erstellung hierarchischer Wissensstrukturen."
    )
    
    prompt = (
        """
        Analyze the following academic text and extract a hierarchical knowledge structure with these components:
        
        1. MAIN TOPIC: The primary subject of the text (a concise phrase, max 5 words)
        
        2. SUBTOPICS: Extract the most important concepts or areas within the main topic. The number of subtopics should naturally reflect the content complexity - use as many as needed to accurately represent the material (typically between 3-10).
        
        3. CHILD SUBTOPICS: For each subtopic, provide 2-4 more specific concepts (these are children of the subtopics)
        
        4. FLASHCARDS: Estimate how many flashcards would be useful (between 10-30)
        
        5. QUESTIONS: Estimate how many test questions would be useful (between 5-15)
        
        6. KEY TERMS: Extract 5-10 important technical terms or concepts with their definitions
        
        7. CONTENT TYPE: Identify the type (e.g., 'lecture', 'textbook', 'scientific paper', 'technical documentation')
        
        Return your analysis in valid JSON format with these keys:
        - main_topic (string)
        - subtopics (array of objects with 'name' and 'child_topics' keys, where 'child_topics' is an array of strings)
        - estimated_flashcards (number)
        - estimated_questions (number)
        - key_terms (array of objects with 'term' and 'definition' keys)
        - content_type (string)
        
        Text to analyze:
        """
        if language != 'de' else
        """
        Analysieren Sie den folgenden akademischen Text und extrahieren Sie eine hierarchische Wissensstruktur mit diesen Komponenten:
        
        1. HAUPTTHEMA: Das primäre Thema des Textes (ein prägnanter Ausdruck, max. 5 Wörter)
        
        2. UNTERTHEMEN: Extrahieren Sie die wichtigsten Konzepte oder Bereiche innerhalb des Hauptthemas. Die Anzahl der Unterthemen sollte die Komplexität des Inhalts natürlich widerspiegeln - verwenden Sie so viele wie nötig, um das Material genau darzustellen (typischerweise zwischen 3-10).
        
        3. UNTERUNTERTHEMEN: Für jedes Unterthema 2-4 spezifischere Konzepte (diese sind Kinder der Unterthemen)
        
        4. KARTEIKARTEN: Schätzen Sie, wie viele Karteikarten nützlich wären (zwischen 10-30)
        
        5. FRAGEN: Schätzen Sie, wie viele Testfragen nützlich wären (zwischen 5-15)
        
        6. SCHLÜSSELBEGRIFFE: Extrahieren Sie 5-10 wichtige Fachbegriffe oder Konzepte mit ihren Definitionen
        
        7. INHALTSTYP: Identifizieren Sie den Typ (z.B. 'Vorlesung', 'Lehrbuch', 'wissenschaftlicher Artikel', 'technische Dokumentation')
        
        Geben Sie Ihre Analyse im gültigen JSON-Format mit diesen Schlüsseln zurück:
        - main_topic (string)
        - subtopics (Array von Objekten mit den Schlüsseln 'name' und 'child_topics', wobei 'child_topics' ein Array von Strings ist)
        - estimated_flashcards (Zahl)
        - estimated_questions (Zahl)
        - key_terms (Array von Objekten mit den Schlüsseln 'term' und 'definition')
        - content_type (string)
        
        Zu analysierender Text:
        """
    )
    
    max_chars = 15000
    if len(text) > max_chars:
        text = text[:max_chars] + "...[text truncated]"
    
    response = query_chatgpt(
        prompt + text, 
        client, 
        system_content=system_content,
        session_id=session_id,
        function_name=function_name
    )
    
    try:
        result = json.loads(response)
        expected_keys = ['main_topic', 'subtopics', 'estimated_flashcards', 'estimated_questions', 'key_terms', 'content_type']
        for key in expected_keys:
            if key not in result:
                if key == 'subtopics':
                    result[key] = []
                elif key == 'key_terms':
                    result[key] = []
                elif key == 'main_topic':
                    result[key] = "Unknown Topic"
                elif key == 'content_type':
                    result[key] = "unknown"
                else:
                    result[key] = 0
        return result
    except json.JSONDecodeError:
        # Fallback when JSON parsing fails
        print("JSON parsing failed, using regex extraction")
        main_topic_match = re.search(r'"main_topic"\s*:\s*"([^"]+)"', response)
        subtopics_match = re.search(r'"subtopics"\s*:\s*\[(.*?)\]', response, re.DOTALL)
        flashcards_match = re.search(r'"estimated_flashcards"\s*:\s*(\d+)', response)
        questions_match = re.search(r'"estimated_questions"\s*:\s*(\d+)', response)
        
        main_topic = main_topic_match.group(1) if main_topic_match else "Unknown Topic"
        subtopics_str = subtopics_match.group(1) if subtopics_match else ""
        
        # Simplified regex-based extraction
        fallback_subtopics = []
        try:
            # Extract objects from the subtopics array
            subtopic_objects = re.findall(r'{[^}]+}', subtopics_str)
            for obj in subtopic_objects:
                name_match = re.search(r'"name"\s*:\s*"([^"]+)"', obj)
                if name_match:
                    name = name_match.group(1)
                    child_topics = []
                    child_topics_match = re.search(r'"child_topics"\s*:\s*\[(.*?)\]', obj, re.DOTALL)
                    if child_topics_match:
                        child_topics_str = child_topics_match.group(1)
                        child_topics = [topic.strip(' "\'') for topic in re.findall(r'"([^"]+)"', child_topics_str)]
                    fallback_subtopics.append({"name": name, "child_topics": child_topics})
        except:
            # If regex extraction fails, use a very simple approach
            subtopics = [s.strip(' "\'') for s in subtopics_str.split(',') if s.strip(' "\'')]
            fallback_subtopics = [{"name": topic, "child_topics": []} for topic in subtopics]
        
        estimated_flashcards = int(flashcards_match.group(1)) if flashcards_match else 15
        estimated_questions = int(questions_match.group(1)) if questions_match else 8
        
        return {
            'main_topic': main_topic,
            'subtopics': fallback_subtopics,
            'estimated_flashcards': estimated_flashcards,
            'estimated_questions': estimated_questions,
            'key_terms': [],
            'content_type': "unknown"
        }

def generate_concept_map_suggestions(text, client, main_topic, parent_subtopics, language='en', analysis_data=None):
    """
    Generiert KI-Vorschläge für Kinder-Subtopics einer Concept-Map.
    Nutzt bereits vorhandene Analysedaten, wenn vorhanden, um doppelte OpenAI-Anfragen zu vermeiden.
    
    Args:
        text: Der Textinhalt für die Generierung
        client: Der OpenAI-Client
        main_topic: Das Hauptthema der Concept-Map
        parent_subtopics: Die übergeordneten Subtopics
        language: Die Sprache der Vorschläge
        analysis_data: Bereits vorhandene Analysedaten aus unified_content_processing
        
    Returns:
        Dictionary mit den übergeordneten Subtopics als Schlüssel und Listen von vorgeschlagenen
        Kinder-Subtopics als Werte, zusammen mit Vorschlägen für Verbindungs-Labels
    """
    # Wenn bereits Analysedaten vorhanden sind, nutze diese direkt
    if analysis_data and 'subtopics' in analysis_data:
        print("Verwende vorhandene Analysedaten für Concept-Map-Vorschläge")
        
        # Ergebnis-Dictionary erstellen
        result = {}
        
        # Subtopics aus Analysedaten extrahieren
        analysis_subtopics = analysis_data['subtopics']
        
        # Für jedes übergeordnete Subtopic
        for parent in parent_subtopics:
            # Suche passendes Subtopic in Analysedaten
            matching_subtopic = None
            for subtopic in analysis_subtopics:
                if isinstance(subtopic, dict) and 'name' in subtopic and subtopic['name'] == parent:
                    matching_subtopic = subtopic
                    break
            
            # Wenn ein passendes Subtopic gefunden wurde und es child_topics hat
            if matching_subtopic and 'child_topics' in matching_subtopic and matching_subtopic['child_topics']:
                # Verwende die vorhandenen child_topics
                result[parent] = matching_subtopic['child_topics']
                
                # Erstelle generische Beziehungslabels
                relationship_labels = {}
                for child in matching_subtopic['child_topics']:
                    relationship_labels[child] = f"is an aspect of {parent}" if language != 'de' else f"ist ein Aspekt von {parent}"
                
                result[f"{parent}_relationships"] = relationship_labels
            else:
                # Fallback für den Fall, dass keine Kind-Themen gefunden wurden
                if language == 'de':
                    children = ["Unterthema 1", "Unterthema 2", "Unterthema 3"]
                    result[parent] = children
                    relationship_labels = {}
                    for child in children:
                        relationship_labels[child] = f"ist ein Aspekt von {parent}"
                    result[f"{parent}_relationships"] = relationship_labels
                else:
                    children = ["Subtopic 1", "Subtopic 2", "Subtopic 3"]
                    result[parent] = children
                    relationship_labels = {}
                    for child in children:
                        relationship_labels[child] = f"is an aspect of {parent}"
                    result[f"{parent}_relationships"] = relationship_labels
        
        return result
    
    # Wenn keine Analysedaten vorhanden sind, führe eine OpenAI-Anfrage durch (bisheriger Code)
    # Importiere benötigte Module
    import json
    import re
    import logging
    import os
    
    system_content = (
        """
        You are an expert in educational knowledge mapping and concept organization. Your task is to analyze text and generate meaningful child subtopics for a concept map.
        
        CRITICAL INSTRUCTIONS:
        - Provide child subtopics that are specific, meaningful, and directly related to their parent topics
        - Ensure each child subtopic represents a distinct and important aspect of its parent topic
        - Use concise, clear terminology that accurately represents the concept
        - Aim for consistency in the level of abstraction across child topics
        - Each child subtopic should be 1-4 words, very concise and focused
        - Provide 2-4 child subtopics for each parent topic
        - For each child subtopic, include a meaningful connection label that explains the relationship between the parent and child topic
        - The connection label should be a short phrase (3-10 words) that clearly describes why these topics are connected
        - ALWAYS use English for all topics and labels
        - RESPOND ONLY WITH VALID JSON - do not include any backticks, markdown formatting, or explanation text
        """
        if language != 'de' else
        """
        Sie sind ein Experte für pädagogische Wissenskartierung und Konzeptorganisation. Ihre Aufgabe ist es, Texte zu analysieren und aussagekräftige Kinder-Unterthemen für eine Concept-Map zu generieren.
        
        KRITISCHE ANWEISUNGEN:
        - Stellen Sie Kinder-Unterthemen bereit, die spezifisch, aussagekräftig und direkt mit ihren übergeordneten Themen verbunden sind
        - Stellen Sie sicher, dass jedes Kinder-Unterthema einen eigenständigen und wichtigen Aspekt seines übergeordneten Themas darstellt
        - Verwenden Sie prägnante, klare Terminologie, die das Konzept genau wiedergibt
        - Streben Sie nach Konsistenz im Abstraktionsgrad über alle Kinderthemen hinweg
        - Jedes Kinder-Unterthema sollte 1-4 Wörter umfassen, sehr prägnant und fokussiert sein
        - Geben Sie 2-4 Kinder-Unterthemen für jedes übergeordnete Thema an
        - Fügen Sie für jedes Kinder-Unterthema ein aussagekräftiges Beziehungslabel an, das die Beziehung zwischen dem Eltern- und dem Kindthema erklärt
        - Das Verbindungslabel sollte eine kurze Phrase (3-10 Wörter) sein, die klar beschreibt, warum diese Themen verbunden sind
        - Verwenden Sie IMMER Deutsch für alle Themen und Labels
        - ANTWORTEN SIE NUR MIT GÜLTIGEM JSON - verwenden Sie keine Backticks, keine Markdown-Formatierung und keine erklärenden Texte
        """
    )
    
    prompt = (
        f"""
        Generate child subtopics for a concept map based on the following structure:
        
        Main Topic: {main_topic}
        
        Parent Subtopics:
        {", ".join(parent_subtopics)}
        
        For each parent subtopic, please suggest 2-4 specific, meaningful child subtopics that represent important aspects of the parent topic. Each child subtopic should be concise (1-4 words).
        
        IMPORTANT: For each child subtopic, also provide a meaningful relationship label that explains why the child topic is connected to its parent topic.
        
        Format your response as a JSON object where each key is a parent subtopic and its value is an array of objects with "topic" and "relationship" properties.
        
        Example format:
        {{
          "Parent Subtopic 1": [
            {{ "topic": "Child Topic 1", "relationship": "is a fundamental component of" }},
            {{ "topic": "Child Topic 2", "relationship": "provides methods for" }},
            {{ "topic": "Child Topic 3", "relationship": "illustrates the application of" }}
          ],
          "Parent Subtopic 2": [
            {{ "topic": "Child Topic 1", "relationship": "emerged historically from" }},
            {{ "topic": "Child Topic 2", "relationship": "represents a specialized case of" }}
          ]
        }}

        IMPORTANT: RESPOND ONLY WITH RAW JSON. DO NOT INCLUDE BACKTICKS OR MARKDOWN FORMATTING.
        """
        if language != 'de' else
        f"""
        Generieren Sie Kinder-Unterthemen für eine Concept-Map basierend auf der folgenden Struktur:
        
        Hauptthema: {main_topic}
        
        Übergeordnete Unterthemen:
        {", ".join(parent_subtopics)}
        
        Schlagen Sie für jedes übergeordnete Unterthema 2-4 spezifische, aussagekräftige Kinder-Unterthemen vor, die wichtige Aspekte des übergeordneten Themas darstellen. Jedes Kinder-Unterthema sollte prägnant sein (1-4 Wörter).
        
        WICHTIG: Geben Sie für jedes Kinder-Unterthema auch ein aussagekräftiges Beziehungslabel an, das erklärt, warum das Kinderthema mit seinem Elternthema verbunden ist.
        
        Formatieren Sie Ihre Antwort als JSON-Objekt, bei dem jeder Schlüssel ein übergeordnetes Unterthema ist und sein Wert ein Array von Objekten mit den Eigenschaften "topic" und "relationship" ist.
        
        Beispielformat:
        {{
          "Übergeordnetes Unterthema 1": [
            {{ "topic": "Kinderthema 1", "relationship": "ist eine grundlegende Komponente von" }},
            {{ "topic": "Kinderthema 2", "relationship": "bietet Methoden für" }},
            {{ "topic": "Kinderthema 3", "relationship": "veranschaulicht die Anwendung von" }}
          ],
          "Übergeordnetes Unterthema 2": [
            {{ "topic": "Kinderthema 1", "relationship": "entwickelte sich historisch aus" }},
            {{ "topic": "Kinderthema 2", "relationship": "repräsentiert einen spezialisierten Fall von" }}
          ]
        }}

        WICHTIG: ANTWORTEN SIE NUR MIT REINEM JSON. KEINE BACKTICKS ODER MARKDOWN-FORMATIERUNG EINFÜGEN.
        """
    )
    
    max_chars = 15000
    if len(text) > max_chars:
        text = text[:max_chars] + "...[text truncated]"
    
    # Caching-System für Concept-Map-Vorschläge
    import hashlib
    cache_key = f"concept_map_{hashlib.md5((prompt + text[:1000]).encode()).hexdigest()}"
    
    # Prüfen, ob Ergebnis im Cache ist
    from flask import current_app
    cache_dir = current_app.config.get("CACHE_DIR", "cache")
    import os
    import json
    
    cache_file = os.path.join(cache_dir, f"{cache_key}.json")
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'r', encoding='utf-8') as f:
                cached_result = json.load(f)
                import logging
                logging.info(f"Cache hit for prompt: {prompt[:100]}...")
                return cached_result
        except Exception as e:
            import logging
            logging.warning(f"Error reading cache: {str(e)}")
    
    # Begrenzte Anzahl von Versuchen, um eine gültige JSON-Antwort zu erhalten
    max_retries = 3
    last_response = None
    
    for attempt in range(max_retries):
        try:
            response = query_chatgpt(prompt + "\n\n" + text, client, system_content, temperature=0.7)
            last_response = response
            
            # Bereinigen der Antwort
            # Entferne alle Markdown-Formatierungen und Backticks
            clean_response = response
            # Entferne Markdown Codeblock-Marker
            clean_response = re.sub(r'```(?:json)?\s*|\s*```', '', clean_response, flags=re.DOTALL)
            # Entferne führende/nachfolgende Leerzeichen
            clean_response = clean_response.strip()
            
            # Versuche, die bereinigte Antwort als JSON zu parsen
            suggestions = json.loads(clean_response)
            
            # Überprüfen, ob das Ergebnis ein Dictionary ist
            if not isinstance(suggestions, dict):
                raise ValueError("Response is not a dictionary")
            
            # Verarbeite die Antwort ins erwartete Format
            result = {}
            for parent, children in suggestions.items():
                # Überprüfe, ob die Kinder im neuen Format (mit topic und relationship) oder im alten Format sind
                if isinstance(children, list) and all(isinstance(child, dict) and "topic" in child for child in children):
                    # Neues Format mit topic und relationship
                    result[parent] = [child["topic"] for child in children]
                    # Speichere die Beziehung im relationship_labels Dictionary
                    relationship_labels = {}
                    for child in children:
                        relationship_labels[child["topic"]] = child.get("relationship", f"is an aspect of {parent}" if language != 'de' else f"ist ein Aspekt von {parent}")
                    result[f"{parent}_relationships"] = relationship_labels
                else:
                    # Altes Format - nur Liste von Themen
                    result[parent] = children
                    # Erstelle generische Beziehungslabels
                    relationship_labels = {}
                    for child in children:
                        relationship_labels[child] = f"is an aspect of {parent}" if language != 'de' else f"ist ein Aspekt von {parent}"
                    result[f"{parent}_relationships"] = relationship_labels
            
            # Validiere, dass alle übergeordneten Subtopics enthalten sind
            for parent in parent_subtopics:
                if parent not in result:
                    if language == 'de':
                        children = ["Beispielkind 1", "Beispielkind 2"]
                        result[parent] = children
                        # Füge generische Beziehungslabels hinzu
                        relationship_labels = {}
                        for child in children:
                            relationship_labels[child] = f"ist ein Aspekt von {parent}"
                        result[f"{parent}_relationships"] = relationship_labels
                    else:
                        children = ["Example Child 1", "Example Child 2"]
                        result[parent] = children
                        # Füge generische Beziehungslabels hinzu
                        relationship_labels = {}
                        for child in children:
                            relationship_labels[child] = f"is an aspect of {parent}"
                        result[f"{parent}_relationships"] = relationship_labels
            
            # Speichern des Ergebnisses im Cache
            try:
                os.makedirs(cache_dir, exist_ok=True)
                with open(cache_file, 'w', encoding='utf-8') as f:
                    json.dump(result, f, ensure_ascii=False, indent=2)
            except Exception as e:
                import logging
                logging.warning(f"Error writing cache: {str(e)}")
            
            return result
            
        except (json.JSONDecodeError, ValueError) as e:
            import logging
            logging.error(f"Error parsing JSON response: {str(e)}")
            logging.error(f"Raw response: {last_response}")
            # Wenn wir den letzten Versuch erreicht haben, brechen wir ab und verwenden die Fallback-Lösung
            if attempt == max_retries - 1:
                break
    
    # Fallback: Erstelle ein einfaches Wörterbuch mit generischen Kindern und Beziehungen
    fallback_suggestions = {}
    for parent in parent_subtopics:
        if language == 'de':
            children = [f"{parent} Komponente 1", f"{parent} Komponente 2", f"{parent} Anwendung"]
            fallback_suggestions[parent] = children
            # Füge generische Beziehungslabels hinzu
            relationship_labels = {}
            for child in children:
                relationship_labels[child] = f"ist ein Bestandteil von {parent}"
            fallback_suggestions[f"{parent}_relationships"] = relationship_labels
        else:
            children = [f"{parent} Component 1", f"{parent} Component 2", f"{parent} Application"]
            fallback_suggestions[parent] = children
            # Füge generische Beziehungslabels hinzu
            relationship_labels = {}
            for child in children:
                relationship_labels[child] = f"is a component of {parent}"
            fallback_suggestions[f"{parent}_relationships"] = relationship_labels
    
    # Speichern des Fallback-Ergebnisses im Cache
    try:
        os.makedirs(cache_dir, exist_ok=True)
        with open(cache_file, 'w', encoding='utf-8') as f:
            json.dump(fallback_suggestions, f, ensure_ascii=False, indent=2)
    except Exception as e:
        import logging
        logging.warning(f"Error writing cache: {str(e)}")
    
    return fallback_suggestions

def generate_additional_flashcards(text, client, analysis, existing_flashcards, num_to_generate=5, language='en', session_id=None, function_name="generate_additional_flashcards"):
    """
    Generiert zusätzliche, einzigartige Flashcards, die sich von den bestehenden unterscheiden.
    
    Args:
        text: Der Textinhalt für die Generierung
        client: Der OpenAI-Client
        analysis: Das Analyseergebnis mit Hauptthema und Unterthemen
        existing_flashcards: Die bereits vorhandenen Flashcards
        num_to_generate: Anzahl der zu generierenden neuen Flashcards
        language: Die Sprache der Flashcards
        session_id: Die ID der aktuellen Session (für Token-Tracking)
        function_name: Name der Funktion (für Token-Tracking)
        
    Returns:
        Liste mit neuen, einzigartigen Flashcards
    """
    # Nutze die vorhandene Funktion, aber mit spezifischen Anweisungen
    system_content = (
        """
        You are an expert in creating ADDITIONAL educational flashcards. Your task is to generate flashcards that are COMPLETELY DIFFERENT from the existing ones.
        
        CRITICAL INSTRUCTIONS:
        - Study the existing flashcards carefully to understand what's already covered
        - Create NEW flashcards that explore DIFFERENT aspects, angles, and applications
        - Do NOT duplicate or rephrase existing flashcards
        - Ensure variety in question types, complexity levels, and topic coverage
        - Focus on areas, concepts, and applications not yet addressed
        - Make each new flashcard substantially different from all existing ones
        - ALWAYS use English for all flashcards
        """
        if language != 'de' else
        """
        Sie sind ein Experte für die Erstellung ZUSÄTZLICHER Lernkarteikarten. Ihre Aufgabe ist es, Karteikarten zu erstellen, die sich VOLLSTÄNDIG von den vorhandenen unterscheiden.
        
        KRITISCHE ANWEISUNGEN:
        - Studieren Sie die vorhandenen Karteikarten sorgfältig, um zu verstehen, was bereits abgedeckt ist
        - Erstellen Sie NEUE Karteikarten, die ANDERE Aspekte, Blickwinkel und Anwendungen erkunden
        - Duplizieren oder umformulieren Sie KEINE vorhandenen Karteikarten
        - Achten Sie auf Vielfalt bei Fragetypen, Komplexitätsstufen und Themenabdeckung
        - Konzentrieren Sie sich auf Bereiche, Konzepte und Anwendungen, die noch nicht behandelt wurden
        - Machen Sie jede neue Karteikarte wesentlich anders als alle vorhandenen
        - Verwenden Sie IMMER Deutsch für alle Karteikarten
        """
    )
    
    prompt = (
        f"""
        I need {num_to_generate} ADDITIONAL, UNIQUE flashcards that are COMPLETELY DIFFERENT from the existing ones.
        
        Main Topic: {analysis.get('main_topic', '')}
        Subtopics: {', '.join([subtopic['name'] if isinstance(subtopic, dict) and 'name' in subtopic else str(subtopic) for subtopic in analysis.get('subtopics', [])])}
        
        EXISTING FLASHCARDS (DO NOT DUPLICATE OR REPHRASE THESE):
        """
        if language != 'de' else
        f"""
        Ich benötige {num_to_generate} ZUSÄTZLICHE, EINZIGARTIGE Karteikarten, die sich VÖLLIG von den vorhandenen unterscheiden.
        
        Hauptthema: {analysis.get('main_topic', '')}
        Unterthemen: {', '.join([subtopic['name'] if isinstance(subtopic, dict) and 'name' in subtopic else str(subtopic) for subtopic in analysis.get('subtopics', [])])}
        
        VORHANDENE KARTEIKARTEN (DIESE NICHT DUPLIZIEREN ODER UMFORMULIEREN):
        """
    )
    
    # Füge die vorhandenen Flashcards zum Prompt hinzu, damit sie nicht dupliziert werden
    for i, fc in enumerate(existing_flashcards[:10]):  # Begrenze auf 10 vorhandene Karten, um den Prompt nicht zu lang zu machen
        if language != 'de':
            prompt += f"\n{i+1}. Q: {fc.get('question', '')} A: {fc.get('answer', '')}"
        else:
            prompt += f"\n{i+1}. F: {fc.get('question', '')} A: {fc.get('answer', '')}"
    
    # Füge Beispiele für die Art von Flashcards hinzu, die wir suchen
    if language != 'de':
        prompt += """
        
        FORMAT REQUIREMENTS:
        - Each flashcard should be in the format "Q: [question] A: [answer]"
        - Questions should be clear, specific, and thought-provoking
        - Answers should be comprehensive but concise, about 2-3 sentences
        - Make the flashcards challenging but still directly relevant to the content
        
        RETURN ONLY THE FLASHCARDS IN THIS JSON FORMAT:
        [
          {"question": "Your question here", "answer": "Your answer here"},
          {"question": "Your question here", "answer": "Your answer here"}
        ]
        """
    else:
        prompt += """
        
        FORMATANFORDERUNGEN:
        - Jede Karteikarte sollte im Format "F: [Frage] A: [Antwort]" sein
        - Fragen sollten klar, spezifisch und zum Nachdenken anregend sein
        - Antworten sollten umfassend aber prägnant sein, etwa 2-3 Sätze
        - Machen Sie die Karteikarten herausfordernd, aber dennoch direkt relevant für den Inhalt
        
        GEBEN SIE NUR DIE KARTEIKARTEN IN DIESEM JSON-FORMAT ZURÜCK:
        [
          {"question": "Ihre Frage hier", "answer": "Ihre Antwort hier"},
          {"question": "Ihre Frage hier", "answer": "Ihre Antwort hier"}
        ]
        """
    
    try:
        # Generiere die neuen Flashcards mit zusätzlichen Parameter für Token-Tracking
        flashcards_response = query_chatgpt(
            prompt, 
            client, 
            system_content=system_content, 
            temperature=0.8,
            session_id=session_id,
            function_name=function_name
        )
        print(f"Flashcards response: {flashcards_response[:200]}...")  # DEBUG: Zeige Anfang der Antwort
        
        # Extrahiere den JSON-Teil aus der Antwort
        match = re.search(r'\[.*\]', flashcards_response, re.DOTALL)
        if match:
            flashcards_json = match.group(0)
            new_flashcards = json.loads(flashcards_json)
            
            # Verifiziere, dass die Antwort im richtigen Format ist
            if isinstance(new_flashcards, list) and all(isinstance(f, dict) and 'question' in f and 'answer' in f for f in new_flashcards):
                return new_flashcards
        
        # Wenn das Format nicht passt oder kein JSON gefunden wurde, versuche erneut zu parsen
        try:
            # Versuche die gesamte Antwort als JSON zu interpretieren
            new_flashcards = json.loads(flashcards_response)
            if isinstance(new_flashcards, list) and all(isinstance(f, dict) and 'question' in f and 'answer' in f for f in new_flashcards):
                return new_flashcards
        except json.JSONDecodeError:
            # Wenn das nicht funktioniert, extrahiere mit einem anderen Regex
            flashcards_list = []
            # Suche nach {"question": ... "answer": ...} Mustern
            pattern = r'{"question":\s*"([^"]+)",\s*"answer":\s*"([^"]+)"}'
            matches = re.findall(pattern, flashcards_response)
            for q, a in matches:
                flashcards_list.append({"question": q, "answer": a})
            
            if flashcards_list:
                return flashcards_list
    except Exception as e:
        print(f"Error generating flashcards: {str(e)}")
    
    # Fallback: Erstelle einfache generische Flashcards, wenn nichts anderes funktioniert
    if language != 'de':
        return [
            {
                "question": f"What are the key components of {analysis.get('main_topic', 'this topic')}?",
                "answer": "The key components include understanding the fundamental principles, practical applications, and theoretical frameworks within this field."
            },
            {
                "question": f"How would you explain {analysis.get('main_topic', 'this concept')} to a beginner?",
                "answer": f"{analysis.get('main_topic', 'This concept')} can be explained as a systematic approach to organizing knowledge and solving problems within a specific domain."
            }
        ]
    else:
        return [
            {
                "question": f"Was sind die Hauptkomponenten von {analysis.get('main_topic', 'diesem Thema')}?",
                "answer": "Die Hauptkomponenten umfassen das Verständnis der grundlegenden Prinzipien, praktischen Anwendungen und theoretischen Rahmenbedingungen in diesem Bereich."
            },
            {
                "question": f"Wie würden Sie {analysis.get('main_topic', 'dieses Konzept')} einem Anfänger erklären?",
                "answer": f"{analysis.get('main_topic', 'Dieses Konzept')} kann als systematischer Ansatz zur Organisation von Wissen und zur Lösung von Problemen in einem bestimmten Bereich erklärt werden."
            }
        ]

def generate_additional_questions(text, client, analysis, existing_questions, num_to_generate=3, language='en', session_id=None, function_name="generate_additional_questions"):
    """
    Generiert zusätzliche, einzigartige Testfragen, die sich von den bestehenden unterscheiden.
    
    Args:
        text: Der Textinhalt für die Generierung
        client: Der OpenAI-Client
        analysis: Das Analyseergebnis mit Hauptthema und Unterthemen
        existing_questions: Die bereits vorhandenen Testfragen
        num_to_generate: Anzahl der zu generierenden neuen Testfragen
        language: Die Sprache der Testfragen
        session_id: Die ID der aktuellen Session (für Token-Tracking)
        function_name: Name der Funktion (für Token-Tracking)
        
    Returns:
        Liste mit neuen, einzigartigen Testfragen
    """
    import time
    import random
    current_time = time.strftime("%Y-%m-%d %H:%M:%S")
    random_seed = random.randint(1000, 9999)
    
    # Nutze die vorhandene Funktion, aber mit spezifischen Anweisungen
    system_content = (
        f"""
        You are an expert in creating ADDITIONAL educational multiple-choice test questions. Your task is to generate questions that are COMPLETELY DIFFERENT from the existing ones.
        
        CRITICAL INSTRUCTIONS:
        - Study the existing questions carefully to understand what's already covered
        - Create NEW questions that explore DIFFERENT aspects, angles, and applications
        - Do NOT duplicate or rephrase existing questions
        - Ensure variety in question types, complexity levels, and topic coverage
        - Focus on areas, concepts, and applications not yet addressed
        - Make each new question substantially different from all existing ones
        - ALWAYS use English for all questions
        - IMPORTANT: Create completely fresh questions (generation time: {current_time}, seed: {random_seed})
        """
        if language != 'de' else
        f"""
        Sie sind ein Experte für die Erstellung ZUSÄTZLICHER Multiple-Choice-Testfragen. Ihre Aufgabe ist es, Fragen zu erstellen, die sich VOLLSTÄNDIG von den vorhandenen unterscheiden.
        
        KRITISCHE ANWEISUNGEN:
        - Studieren Sie die vorhandenen Fragen sorgfältig, um zu verstehen, was bereits abgedeckt ist
        - Erstellen Sie NEUE Fragen, die ANDERE Aspekte, Blickwinkel und Anwendungen erkunden
        - Duplizieren oder umformulieren Sie KEINE vorhandenen Fragen
        - Achten Sie auf Vielfalt bei Fragetypen, Komplexitätsstufen und Themenabdeckung
        - Konzentrieren Sie sich auf Bereiche, Konzepte und Anwendungen, die noch nicht behandelt wurden
        - Machen Sie jede neue Frage wesentlich anders als alle vorhandenen
        - Verwenden Sie IMMER Deutsch für alle Fragen
        - WICHTIG: Erstellen Sie völlig neue Fragen (Generierungszeit: {current_time}, Seed: {random_seed})
        """
    )
    
    prompt = (
        f"""
        I need {num_to_generate} ADDITIONAL, UNIQUE multiple-choice test questions that are COMPLETELY DIFFERENT from the existing ones.
        
        Main Topic: {analysis.get('main_topic', '')}
        Subtopics: {', '.join([subtopic['name'] if isinstance(subtopic, dict) and 'name' in subtopic else str(subtopic) for subtopic in analysis.get('subtopics', [])])}
        
        EXISTING QUESTIONS (DO NOT DUPLICATE OR REPHRASE THESE):
        """
        if language != 'de' else
        f"""
        Ich benötige {num_to_generate} ZUSÄTZLICHE, EINZIGARTIGE Multiple-Choice-Testfragen, die sich VÖLLIG von den vorhandenen unterscheiden.
        
        Hauptthema: {analysis.get('main_topic', '')}
        Unterthemen: {', '.join([subtopic['name'] if isinstance(subtopic, dict) and 'name' in subtopic else str(subtopic) for subtopic in analysis.get('subtopics', [])])}
        
        VORHANDENE FRAGEN (DIESE NICHT DUPLIZIEREN ODER UMFORMULIEREN):
        """
    )
    
    # Füge die vorhandenen Fragen zum Prompt hinzu, damit sie nicht dupliziert werden
    for i, q in enumerate(existing_questions[:5]):  # Begrenze auf 5 vorhandene Fragen, um den Prompt nicht zu lang zu machen
        prompt += f"\n{i+1}. Q: {q.get('text', '')}"
        if 'options' in q:
            prompt += f"\nOptions: {q.get('options', [])}"
        if 'correct' in q:
            prompt += f"\nCorrect Answer: {q.get('correct', 0)}"
        if 'explanation' in q:
            prompt += f"\nExplanation: {q.get('explanation', '')}"
    
    # Füge Beispiele für die Art von Testfragen hinzu, die wir suchen
    if language != 'de':
        prompt += """
        
        FORMAT REQUIREMENTS:
        - Each question should have a clear, specific prompt
        - Provide exactly 4 options for each question
        - Exactly one option should be correct
        - Include a brief explanation of why the correct answer is right
        - Vary question types (e.g., "which is NOT...", "what is the best...", etc.)
        
        RETURN ONLY THE QUESTIONS IN THIS JSON FORMAT WITHOUT ANY OTHER TEXT:
        [
          {
            "text": "Your question here",
            "options": ["Option A", "Option B", "Option C", "Option D"],
            "correct": 0,  // Index of correct option (0-3)
            "explanation": "Brief explanation of the correct answer"
          }
        ]
        """
    else:
        prompt += """
        
        FORMATANFORDERUNGEN:
        - Jede Frage sollte eine klare, spezifische Formulierung haben
        - Geben Sie genau 4 Optionen für jede Frage an
        - Genau eine Option sollte korrekt sein
        - Fügen Sie eine kurze Erläuterung bei, warum die richtige Antwort korrekt ist
        - Variieren Sie die Fragetypen (z.B. "welches ist NICHT...", "was ist am besten...", usw.)
        
        GEBEN SIE NUR DIE FRAGEN IN DIESEM JSON-FORMAT OHNE ZUSÄTZLICHEN TEXT ZURÜCK:
        [
          {
            "text": "Ihre Frage hier",
            "options": ["Option A", "Option B", "Option C", "Option D"],
            "correct": 0,  // Index der richtigen Option (0-3)
            "explanation": "Kurze Erklärung der richtigen Antwort"
          }
        ]
        """
    
    try:
        # Generiere die neuen Testfragen mit Token-Tracking-Parametern
        questions_response = query_chatgpt(
            prompt, 
            client, 
            system_content=system_content, 
            temperature=0.9,  # Erhöhte Temperatur für mehr Variabilität
            session_id=session_id,
            function_name=function_name
        )
        
        print(f"Questions response: {questions_response[:200]}...")  # DEBUG: Zeige Anfang der Antwort
        
        # Mehrere Strategien zum Extrahieren der JSON-Daten
        
        # Strategie 1: Direkte JSON-Konvertierung der gesamten Antwort probieren
        try:
            new_questions = json.loads(questions_response)
            if isinstance(new_questions, list) and all(isinstance(q, dict) and 'text' in q and 'options' in q and 'correct' in q for q in new_questions):
                print("Strategie 1 erfolgreich: JSON direkt geparst")
                return new_questions
        except json.JSONDecodeError:
            print("Strategie 1 fehlgeschlagen: Konnte Antwort nicht direkt als JSON parsen")
            pass
        
        # Strategie 2: Suche nach JSON-Array mit Regex
        try:
            match = re.search(r'\[\s*\{.*\}\s*\]', questions_response, re.DOTALL)
            if match:
                json_text = match.group(0)
                new_questions = json.loads(json_text)
                if isinstance(new_questions, list) and all(isinstance(q, dict) and 'text' in q and 'options' in q and 'correct' in q for q in new_questions):
                    print("Strategie 2 erfolgreich: JSON mit Regex extrahiert")
                    return new_questions
        except json.JSONDecodeError:
            print("Strategie 2 fehlgeschlagen: Konnte extrahierten Array-Text nicht als JSON parsen")
            pass
        
        # Strategie 3: Bereinige Markdown-Code-Blöcke
        try:
            # Entferne Markdown-Code-Block-Marker und Sprachhinweise
            cleaned_response = re.sub(r'```(?:json)?\s*|\s*```', '', questions_response)
            new_questions = json.loads(cleaned_response)
            if isinstance(new_questions, list) and all(isinstance(q, dict) and 'text' in q and 'options' in q and 'correct' in q for q in new_questions):
                print("Strategie 3 erfolgreich: Markdown-Blöcke entfernt und als JSON geparst")
                return new_questions
        except json.JSONDecodeError:
            print("Strategie 3 fehlgeschlagen: Konnte bereinigten Text nicht als JSON parsen")
            pass
        
        # Strategie 4: Einzelne JSON-Objekte extrahieren
        try:
            questions_list = []
            pattern = r'{\s*"text"\s*:\s*"([^"]*)"\s*,\s*"options"\s*:\s*(\[[^\]]*\])\s*,\s*"correct"\s*:\s*(\d+)\s*,\s*"explanation"\s*:\s*"([^"]*)"\s*}'
            matches = re.findall(pattern, questions_response)
            
            for text, options_str, correct, explanation in matches:
                try:
                    # Bereinige mögliche Escape-Sequenzen in den Optionen
                    cleaned_options = options_str.replace('\\', '\\\\').replace('\\"', '\\\\"')
                    options = json.loads(cleaned_options)
                    questions_list.append({
                        "text": text,
                        "options": options,
                        "correct": int(correct),
                        "explanation": explanation
                    })
                except Exception as e:
                    print(f"Fehler beim Parsen einer einzelnen Frage: {str(e)}")
                    continue
            
            if questions_list and len(questions_list) >= min(1, num_to_generate // 2):
                print(f"Strategie 4 erfolgreich: {len(questions_list)} individuelle Fragen extrahiert")
                return questions_list
        except Exception as e:
            print(f"Strategie 4 fehlgeschlagen: {str(e)}")
            pass
        
        # Wenn alle Strategien fehlschlagen, gebe einen Hinweis im Log und verwende den Fallback
        print("Alle Parsing-Strategien fehlgeschlagen. Text der Antwort:")
        print(questions_response)
        
    except Exception as e:
        print(f"Error generating questions: {str(e)}")
    
    # Fallback: Erstelle einfache generische Testfragen, wenn nichts anderes funktioniert
    print("Verwende Fallback-Fragen")
    if language != 'de':
        return [
            {
                "text": f"Which of the following best describes {analysis.get('main_topic', 'this topic')}?",
                "options": [
                    "A systematic approach to problem-solving in the field",
                    "A random collection of unrelated concepts",
                    "A purely theoretical framework with no practical applications",
                    "A deprecated methodology no longer used in modern practice"
                ],
                "correct": 0,
                "explanation": f"{analysis.get('main_topic', 'This topic')} is primarily characterized by its systematic approach to problem-solving."
            },
            {
                "text": f"Which of the following is NOT a key component of {analysis.get('main_topic', 'this field')}?",
                "options": [
                    "Theoretical foundations",
                    "Practical applications",
                    "Systematic methodology",
                    "Arbitrary decision-making"
                ],
                "correct": 3,
                "explanation": "Arbitrary decision-making contradicts the systematic nature of this field, which relies on structured approaches."
            }
        ]
    else:
        return [
            {
                "text": f"Welche der folgenden Beschreibungen trifft am besten auf {analysis.get('main_topic', 'dieses Thema')} zu?",
                "options": [
                    "Ein systematischer Ansatz zur Problemlösung in diesem Bereich",
                    "Eine zufällige Sammlung unzusammenhängender Konzepte",
                    "Ein rein theoretisches Rahmenwerk ohne praktische Anwendungen",
                    "Eine veraltete Methodik, die in der modernen Praxis nicht mehr verwendet wird"
                ],
                "correct": 0,
                "explanation": f"{analysis.get('main_topic', 'Dieses Thema')} zeichnet sich hauptsächlich durch seinen systematischen Ansatz zur Problemlösung aus."
            },
            {
                "text": f"Welche der folgenden Komponenten ist KEINE Schlüsselkomponente von {analysis.get('main_topic', 'diesem Bereich')}?",
                "options": [
                    "Theoretische Grundlagen",
                    "Praktische Anwendungen",
                    "Systematische Methodik",
                    "Willkürliche Entscheidungsfindung"
                ],
                "correct": 3,
                "explanation": "Willkürliche Entscheidungsfindung widerspricht der systematischen Natur dieses Bereichs, der auf strukturierten Ansätzen basiert."
            }
        ]

def unified_content_processing(text, client, file_names=None, user_id=None, language=None, max_retries=3, session_id=None):
    """
    Verarbeitet den Text und generiert alle notwendigen Inhalte (Analyse, Flashcards, Fragen) in einem einzigen API-Aufruf.
    
    Args:
        text: Der zu verarbeitende Text
        client: Der OpenAI-Client oder OptimizedOpenAIClient-Klasse
        file_names: Liste der Dateinamen (für intelligenten Fallback)
        user_id: Die ID des Benutzers (für Celery Tasks, wo g.user nicht verfügbar ist)
        language: Die Sprache des Textes (wird automatisch erkannt, wenn nicht angegeben)
        max_retries: Maximale Anzahl von Wiederholungsversuchen
        session_id: Die ID der Session für das Logging
    
    Returns:
        dict: Die kombinierte Antwort mit allen Komponenten
    
    Raises:
        ValueError: Wenn der Client fehlt oder falsch konfiguriert ist
    """
    import logging
    import time
    import traceback
    import json
    import os
    import sys
    from api.log_utils import AppLogger
    from api.openai_client import OptimizedOpenAIClient
    
    logger = logging.getLogger(__name__)
    
    # DEBUG-Logging mit strukturiertem Logger
    AppLogger.structured_log(
        "INFO",
        f"unified_content_processing für Session {session_id} gestartet",
        session_id=session_id,
        component="content_processor",
        text_length=len(text) if text else 0,
        files_count=len(file_names) if file_names else 0
    )
    
    # Prüfen, ob der Client eine Klasse oder eine Instanz ist
    is_optimized_client = client == OptimizedOpenAIClient
    
    # Validierung für den Client - unterschiedlich je nach Typ
    if not is_optimized_client and client is None:
        error_msg = "OpenAI-Client ist None - kann API-Anfrage nicht durchführen"
        AppLogger.track_error(
            session_id, 
            "client_initialization_error", 
            error_msg
        )
        raise ValueError(error_msg)
    
    if language is None:
        language = detect_language(text)
        
    AppLogger.structured_log(
        "INFO",
        f"Erkannte Sprache: {language}",
        session_id=session_id,
        component="language_detection"
    )
    
    # Erhöhe das Zeichenlimit deutlich für GPT-4o
    # Verwendet etwa 80% der verfügbaren Tokens für den Eingabetext
    max_chars = 380000  # Etwa 95K Tokens für den Eingabetext
    
    if len(text) > max_chars:
        AppLogger.structured_log(
            "WARNING",
            f"Text zu lang: {len(text)} Zeichen, kürze auf {max_chars} Zeichen",
            session_id=session_id,
            component="content_processor"
        )
        text = text[:max_chars] + "...[Text abgeschnitten]"
    
    # Verwende jetzt die präzise Tokenzählung mit tiktoken
    from api.token_tracking import count_tokens
    document_tokens = count_tokens(text)
    AppLogger.structured_log(
        "INFO",
        f"Textanalyse: {len(text)} Zeichen, {document_tokens} Tokens",
        session_id=session_id,
        component="token_counter",
        tokens=document_tokens
    )
    
    # Anzahl der Dokumente ermitteln für dynamische Anpassung der Anzahl von Flashcards und Fragen
    num_documents = len(file_names) if file_names else 1
    
    # Berechne die Mindestanzahl von Flashcards und Fragen basierend auf der Dokumentanzahl
    min_flashcards = max(8, num_documents * 5)    # Mindestens 5 pro Dokument, aber nicht weniger als 8 gesamt
    max_flashcards = min(30, min_flashcards * 2)  # Maximal 30, aber flexibel nach oben
    
    min_questions = max(5, num_documents * 3)     # Mindestens 3 pro Dokument, aber nicht weniger als 5 gesamt
    max_questions = min(20, min_questions * 2)    # Maximal 20, aber flexibel nach oben
    
    AppLogger.structured_log(
        "INFO",
        f"Zielwerte: {min_flashcards}-{max_flashcards} Flashcards, {min_questions}-{max_questions} Fragen",
        session_id=session_id,
        component="content_processor",
        documents=num_documents
    )
    
    # System-Prompt für einen sehr spezialisierten Assistenten
    system_content = (
        """
        You are an AI academic assistant specialized in analyzing educational content and creating structured learning materials.
        Your task is to process academic text and generate a complete set of learning materials including:
        1. Hierarchical topic structure
        2. Flashcards for studying
        3. Multiple-choice questions for testing knowledge
        
        IMPORTANT: The input text might contain content from multiple documents combined together.
        Carefully analyze ALL of the provided content, including all documents, to create comprehensive learning materials.
        
        Follow these guidelines precisely:
        - Be comprehensive yet concise
        - Process and incorporate content from ALL provided documents
        - Focus on factual information across the entire corpus
        - Ensure all materials are self-contained and don't reference external resources
        - Create original, diverse content that explores topics from multiple angles
        - Do NOT include references to visual elements (figures, tables, etc.)
        - Ensure JSON formatting is valid and follows the specified structure exactly
        - IMPORTANT: Extract as many subtopics as appropriate for the content - do NOT limit to any fixed number
        """
        if language != 'de' else
        """
        Sie sind ein KI-Akademiker-Assistent, der auf die Analyse von Bildungsinhalten und die Erstellung strukturierter Lernmaterialien spezialisiert ist.
        Ihre Aufgabe ist es, akademischen Text zu verarbeiten und einen vollständigen Satz von Lernmaterialien zu erstellen, darunter:
        1. Hierarchische Themenstruktur
        2. Karteikarten zum Lernen
        3. Multiple-Choice-Fragen zum Testen des Wissens
        
        WICHTIG: Der Eingabetext kann Inhalte aus mehreren zusammengefügten Dokumenten enthalten.
        Analysieren Sie sorgfältig ALLE bereitgestellten Inhalte, einschliesslich aller Dokumente, um umfassende Lernmaterialien zu erstellen.
        
        Befolgen Sie diese Richtlinien genau:
        - Umfassend, aber prägnant
        - Verarbeiten und integrieren Sie Inhalte aus ALLEN bereitgestellten Dokumenten
        - Konzentration auf sachliche Informationen aus dem gesamten Textkorpus
        - Stellen Sie sicher, dass alle Materialien eigenständig sind und nicht auf externe Ressourcen verweisen
        - Erstellen Sie originale, vielfältige Inhalte, die Themen aus verschiedenen Blickwinkeln betrachten
        - Nehmen Sie KEINE Verweise auf visuelle Elemente auf (Abbildungen, Tabellen usw.)
        - Stellen Sie sicher, dass die JSON-Formatierung gültig ist und genau der angegebenen Struktur folgt
        - WICHTIG: Extrahieren Sie so viele Unterthemen wie für den Inhalt angemessen - beschränken Sie sich NICHT auf eine feste Anzahl
        """
    )
    
    AppLogger.structured_log(
        "INFO",
        "System-Prompt erstellt",
        session_id=session_id,
        component="prompt_creation",
        prompt_length=len(system_content)
    )
    
    # Implementierung des restlichen Codes für die Funktion
    # Platzhalter, um die Funktion syntaktisch korrekt zu beenden
    return {
        "main_topic": "Automatisch generiert",
        "subtopics": [],
        "estimated_flashcards": min_flashcards,
        "estimated_questions": min_questions,
        "key_terms": [],
        "content_type": "unknown"
    }

def check_and_manage_user_sessions(user_id, max_sessions=5, session_to_exclude=None):
    """
    Überprüft und begrenzt die Anzahl der aktiven Sessions eines Benutzers.
    Löscht die ältesten Sessions (basierend auf last_used_at), wenn die maximale Anzahl überschritten wird.
    NULL-Werte in last_used_at werden als älteste betrachtet und priorisiert gelöscht.
    
    Args:
        user_id (str): Die ID des Benutzers
        max_sessions (int): Maximale Anzahl an Sessions, die behalten werden sollen (Standard: 5)
        session_to_exclude (str): Eine Session-ID, die von der Zählung ausgeschlossen werden soll
    
    Returns:
        bool: True, wenn Sessions gelöscht wurden, False wenn keine Sessions gelöscht wurden
    """
    if not user_id:
        return False
    
    sessions_removed = False
    
    # SQL-Abfrage, um Uploads zu holen:
    # 1. Sessions mit NULL in last_used_at zuerst
    # 2. Dann nach last_used_at aufsteigend sortiert (älteste zuerst)
    user_sessions = Upload.query.filter_by(
        user_id=user_id
    ).order_by(
        Upload.last_used_at.is_(None).desc(),  # NULL-Werte zuerst
        Upload.last_used_at.asc()              # Dann nach Alter sortiert (älteste zuerst)
    ).all()
    
    # Wenn session_to_exclude angegeben ist, diese ausschließen
    if session_to_exclude:
        user_sessions = [s for s in user_sessions if s.session_id != session_to_exclude]
    
    # Anzahl der zu löschenden Sessions berechnen
    total_sessions = len(user_sessions)
    
    # Wenn die Anzahl der Sessions das Maximum überschreitet, lösche die ältesten
    if total_sessions > max_sessions:
        # Berechne, wie viele Sessions gelöscht werden müssen
        sessions_to_delete = total_sessions - max_sessions
        
        AppLogger.structured_log(
            "INFO",
            f"Benutzer {user_id} hat {total_sessions} Sessions, Maximum ist {max_sessions}. Lösche {sessions_to_delete} älteste Session(s).",
            user_id=user_id,
            component="check_and_manage_user_sessions"
        )
        
        # Lösche die ältesten Sessions (die ersten in der Liste)
        for i in range(sessions_to_delete):
            if i < len(user_sessions):
                session_to_remove = user_sessions[i]
                
                # Debug-Ausgabe für last_used_at-Wert
                last_used_value = "NULL" if session_to_remove.last_used_at is None else session_to_remove.last_used_at.isoformat()
                AppLogger.structured_log(
                    "INFO",
                    f"Lösche Session {session_to_remove.session_id} mit last_used_at={last_used_value}",
                    user_id=user_id,
                    component="check_and_manage_user_sessions"
                )
                
                # Lösche alle verknüpften Daten (Themen, Karteikarten, Fragen, etc.)
                try:
                    # Lösche zugehörige Karteikarten
                    Flashcard.query.filter_by(upload_id=session_to_remove.id).delete()
                    
                    # Lösche zugehörige Themen
                    Topic.query.filter_by(upload_id=session_to_remove.id).delete()
                    
                    # Lösche zugehörige Fragen
                    Question.query.filter_by(upload_id=session_to_remove.id).delete()
                    
                    # Lösche zugehörige Verbindungen
                    Connection.query.filter_by(upload_id=session_to_remove.id).delete()
                    
                    # Lösche zugehörige Benutzeraktivitäten
                    UserActivity.query.filter_by(session_id=session_to_remove.session_id).delete()
                    
                    # Lösche den Upload-Eintrag selbst
                    db.session.delete(session_to_remove)
                    
                    # Commit nach jedem gelöschten Upload
                    db.session.commit()
                    
                    sessions_removed = True
                    AppLogger.structured_log(
                        "INFO",
                        f"Upload {session_to_remove.session_id} gelöscht, damit nur die {max_sessions} neuesten Sessions erhalten bleiben",
                        user_id=user_id,
                        session_id=session_to_remove.session_id,
                        component="check_and_manage_user_sessions"
                    )
                    
                    # Lösche auch alle in Redis gespeicherten Daten für diese Session
                    try:
                        from redis_connection import redis_client
                        keys_to_delete = [
                            f"processing_status:{session_to_remove.session_id}",
                            f"processing_progress:{session_to_remove.session_id}",
                            f"processing_start_time:{session_to_remove.session_id}",
                            f"processing_heartbeat:{session_to_remove.session_id}",
                            f"processing_last_update:{session_to_remove.session_id}",
                            f"processing_details:{session_to_remove.session_id}",
                            f"processing_result:{session_to_remove.session_id}",
                            f"task_id:{session_to_remove.session_id}",
                            f"error_details:{session_to_remove.session_id}",
                            f"openai_error:{session_to_remove.session_id}",
                            f"all_data_stored:{session_to_remove.session_id}",
                            f"finalization_complete:{session_to_remove.session_id}"
                        ]
                        
                        # Überprüfe, ob redis_client existiert und initialisiert ist
                        if redis_client:
                            try:
                                # Verwende Redis pipeline für effizientes Löschen
                                pipeline = redis_client.pipeline()
                                for key in keys_to_delete:
                                    pipeline.delete(key)
                                pipeline.execute()
                                
                                AppLogger.structured_log(
                                    "INFO",
                                    f"Redis-Daten für Session {session_to_remove.session_id} gelöscht",
                                    user_id=user_id,
                                    session_id=session_to_remove.session_id,
                                    component="check_and_manage_user_sessions"
                                )
                            except Exception as redis_error:
                                AppLogger.structured_log(
                                    "ERROR",
                                    f"Fehler beim Löschen der Redis-Daten: {str(redis_error)}",
                                    user_id=user_id,
                                    session_id=session_to_remove.session_id,
                                    component="check_and_manage_user_sessions",
                                    exception=traceback.format_exc()
                                )
                        else:
                            AppLogger.structured_log(
                                "WARNING",
                                f"Redis-Client nicht verfügbar, Redis-Daten für Session {session_to_remove.session_id} konnten nicht gelöscht werden",
                                user_id=user_id,
                                session_id=session_to_remove.session_id,
                                component="check_and_manage_user_sessions"
                            )
                    except Exception as e:
                        db.session.rollback()
                        AppLogger.structured_log(
                            "ERROR",
                            f"Fehler beim Löschen des Uploads: {str(e)}",
                            user_id=user_id,
                            session_id=session_to_remove.session_id,
                            component="check_and_manage_user_sessions",
                            exception=traceback.format_exc()
                        )
    
    # Gibt TRUE zurück, wenn Sessions gelöscht wurden
    return sessions_removed

def update_session_timestamp(session_id):
    """
    Aktualisiert den last_used_at-Timestamp einer Session auf die aktuelle Zeit.
    
    Args:
        session_id (str): Die ID der Session
        
    Returns:
        bool: True bei Erfolg, False bei Fehler
    """
    if not session_id:
        return False
    
    try:
        # Finde die Session
        upload = Upload.query.filter_by(session_id=session_id).first()
        if not upload:
            return False
        
        # Setze den Zeitstempel auf die aktuelle Zeit
        upload.last_used_at = db.func.current_timestamp()
        
        # Aktualisiere auch den processing_status, falls nötig
        if upload.processing_status in ['pending', 'error', 'failed']:
            upload.processing_status = 'completed'
        
        db.session.commit()
        
        # Protokolliere die Aktualisierung
        AppLogger.structured_log(
            "INFO",
            f"Zeitstempel für Session {session_id} aktualisiert",
            component="session_management",
            user_id=upload.user_id,
            session_id=session_id
        )
        
        return True
    except Exception as e:
        AppLogger.track_error(
            session_id,
            "session_timestamp_error",
            f"Fehler beim Aktualisieren des Zeitstempels für Session {session_id}: {str(e)}",
            trace=traceback.format_exc()
        )
        db.session.rollback()
        
    return False
