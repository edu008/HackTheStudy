"""
Modul für Textextraktion aus verschiedenen Dokumenttypen.
"""
import os
import logging
import json
import traceback
from typing import Dict, List, Tuple, Any, Optional

# Logger konfigurieren
logger = logging.getLogger(__name__)

def extract_text_from_file(file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Extrahiert Text aus einer Datei basierend auf dem Dateityp.
    
    Args:
        file_path: Pfad zur Datei
        
    Returns:
        Tuple mit (extrahiertem Text, Liste von Chunks mit Metadaten)
    """
    logger.info("="*50)
    logger.info(f"TEXTEXTRAKTION GESTARTET: {file_path}")
    
    file_extension = os.path.splitext(file_path)[1].lower()
    
    if file_extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == '.docx':
        return extract_text_from_docx(file_path)
    elif file_extension == '.txt':
        return extract_text_from_txt(file_path)
    else:
        logger.warning(f"Unbekannter Dateityp: {file_extension}, Fallback auf einfaches Textlesen")
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                text = f.read()
            return text, [{"page": 1, "text": text}]
        except Exception as e:
            logger.error(f"Fehler beim Lesen der Datei: {e}")
            return "", []

def extract_text_from_pdf(file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Extrahiert Text aus einer PDF-Datei mit PyMuPDF (fitz).
    
    Args:
        file_path: Pfad zur PDF-Datei
        
    Returns:
        Tuple mit (extrahiertem Text, Liste von Chunks mit Metadaten)
    """
    try:
        # Importiere PyMuPDF (fitz)
        import fitz
        logger.info(f"[PDF] Extrahiere Text aus PDF: {file_path}")
        
        # Öffne PDF-Dokument
        pdf_document = fitz.open(file_path)
        num_pages = len(pdf_document)
        logger.info(f"[PDF] PDF hat {num_pages} Seiten")
        
        all_text = ""
        chunks = []
        
        # Extrahiere Text seitenweise
        for page_num in range(num_pages):
            page = pdf_document[page_num]
            
            # Text extrahieren
            page_text = page.get_text()
            
            # Text zum Gesamttext hinzufügen
            all_text += page_text + "\n\n"
            
            # Chunk mit Metadaten hinzufügen
            chunks.append({
                "page": page_num + 1,
                "text": page_text,
                "page_width": page.rect.width,
                "page_height": page.rect.height
            })
            
            # Fortschritt loggen bei größeren PDFs
            if num_pages > 20 and page_num % 10 == 0 and page_num > 0:
                logger.debug(f"[PDF] {page_num}/{num_pages} Seiten verarbeitet...")
        
        # PDF schließen
        pdf_document.close()
        
        logger.info(f"[PDF] PDF-Extraktion abgeschlossen: {num_pages} Seiten, {len(all_text)} Zeichen")
        
        # Erste und letzte Seite als Beispiel loggen
        if chunks:
            logger.debug(f"[PDF] Erste Seite (Vorschau): {chunks[0]['text'][:150]}...")
            if len(chunks) > 1:
                logger.debug(f"[PDF] Letzte Seite (Vorschau): {chunks[-1]['text'][:150]}...")
        
        logger.info("="*50)
        return all_text, chunks
    
    except ImportError as e:
        logger.error(f"[PDF] PyMuPDF (fitz) nicht installiert: {e}")
        logger.error("[PDF] Bitte installieren Sie PyMuPDF: pip install pymupdf")
        return "", []
    
    except Exception as e:
        logger.error(f"[PDF] Fehler bei der PDF-Extraktion: {e}")
        logger.error(traceback.format_exc())
        return "", []

def extract_text_from_docx(file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Extrahiert Text aus einer DOCX-Datei.
    
    Args:
        file_path: Pfad zur DOCX-Datei
        
    Returns:
        Tuple mit (extrahiertem Text, Liste von Chunks mit Metadaten)
    """
    try:
        # python-docx für DOCX-Dateien
        import docx
        logger.info(f"[DOCX] Extrahiere Text aus DOCX: {file_path}")
        
        doc = docx.Document(file_path)
        
        # Extrahiere Text aus Paragraphen
        all_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        
        chunks = [{
            "paragraph_index": i,
            "text": para.text
        } for i, para in enumerate(doc.paragraphs) if para.text.strip()]
        
        logger.info(f"[DOCX] DOCX-Extraktion abgeschlossen: {len(chunks)} Paragraphen, {len(all_text)} Zeichen")
        logger.info("="*50)
        return all_text, chunks
    
    except ImportError as e:
        logger.error(f"[DOCX] python-docx nicht installiert: {e}")
        logger.error("[DOCX] Bitte installieren Sie python-docx: pip install python-docx")
        return "", []
    
    except Exception as e:
        logger.error(f"[DOCX] Fehler bei der DOCX-Extraktion: {e}")
        logger.error(traceback.format_exc())
        return "", []

def extract_text_from_txt(file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
    """
    Extrahiert Text aus einer Textdatei.
    
    Args:
        file_path: Pfad zur Textdatei
        
    Returns:
        Tuple mit (extrahiertem Text, Liste von Chunks mit Metadaten)
    """
    try:
        logger.info(f"[TXT] Extrahiere Text aus Textdatei: {file_path}")
        
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read()
        
        # Teile Text in Absätze
        paragraphs = [p for p in text.split('\n\n') if p.strip()]
        
        chunks = [{
            "paragraph_index": i,
            "text": para
        } for i, para in enumerate(paragraphs) if para.strip()]
        
        logger.info(f"[TXT] TXT-Extraktion abgeschlossen: {len(chunks)} Absätze, {len(text)} Zeichen")
        logger.info("="*50)
        return text, chunks
    
    except Exception as e:
        logger.error(f"[TXT] Fehler bei der TXT-Extraktion: {e}")
        logger.error(traceback.format_exc())
        return "", []

# Exportiere die Funktionen
__all__ = [
    'extract_text_from_pdf',
    'extract_text_from_docx',
    'extract_text_from_txt',
    'extract_text_from_file'
] 